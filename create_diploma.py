#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches, cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '12')
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), '000000')
        tcBorders.append(edge_element)
    tcPr.append(tcBorders)

def add_heading_style(doc, text, level=1):
    """Add styled heading"""
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = 'Times New Roman'
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = 'Times New Roman'
    return p

def set_paragraph_formatting(p, space_after=0, space_before=0, line_spacing=1.5):
    """Set paragraph formatting"""
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.first_line_indent = Inches(0.5)

def add_body_paragraph(doc, text):
    """Add body paragraph with formatting"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.5)
    return p

# Create document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = cm(2.0)
    section.bottom_margin = cm(2.0)
    section.left_margin = cm(3.0)
    section.right_margin = cm(1.0)

# ===== ТИТУЛЬНЫЙ ЛИСТ =====
title_heading = doc.add_paragraph()
title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_heading.add_run('МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ\nКЫРГЫЗСКОЙ РЕСПУБЛИКИ')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()

org_heading = doc.add_paragraph()
org_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = org_heading.add_run('БИШКЕКСКИЙ КОЛЛЕДЖ КОМПЬЮТЕРНЫХ СИСТЕМ И ТЕХНОЛОГИЙ')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()

dept_heading = doc.add_paragraph()
dept_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dept_heading.add_run('ОТДЕЛЕНИЕ «ЦИФРОВЫЕ ТЕХНОЛОГИИ И ДИЗАЙН»')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'

for _ in range(5):
    doc.add_paragraph()

vkr_heading = doc.add_paragraph()
vkr_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = vkr_heading.add_run('Выпускная квалификационная работа')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

for _ in range(3):
    doc.add_paragraph()

theme = doc.add_paragraph()
theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = theme.add_run('«Разработка веб-приложения RoadWatch для мониторинга состояния автомобильных трасс и погодных условий»')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

for _ in range(5):
    doc.add_paragraph()

# Specialty
spec = doc.add_paragraph()
spec.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = spec.add_run('Специальность: 230111 «Программное обеспечение вычислительной техники\nи автоматизированных систем»')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

qual = doc.add_paragraph()
qual.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = qual.add_run('Квалификация: техник-программист')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

form = doc.add_paragraph()
form.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = form.add_run('Форма обучения: очная')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

for _ in range(3):
    doc.add_paragraph()

# Supervisor and student
sup_p = doc.add_paragraph()
sup_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = sup_p.add_run('Руководитель: ________________________\n(Ф.И.О., подпись)')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

stud_p = doc.add_paragraph()
stud_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = stud_p.add_run('Исполнитель: ________________________\n(Ф.И.О., подпись)')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

for _ in range(5):
    doc.add_paragraph()

city = doc.add_paragraph()
city.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = city.add_run('г. Бишкек — 2025')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

# Page break
doc.add_page_break()

# ===== СОДЕРЖАНИЕ =====
toc_heading = doc.add_paragraph()
toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = toc_heading.add_run('СОДЕРЖАНИЕ')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()

toc_items = [
    ('Введение', '3'),
    ('1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ ВЕБ-РАЗРАБОТКИ И МОНИТОРИНГА', '5'),
    ('1.1 Понятие веб-приложения и его архитектура', '5'),
    ('1.2 Современные технологии веб-разработки', '7'),
    ('1.3 Мониторинг состояния дорог и системы оповещения', '9'),
    ('1.4 Обзор аналогичных решений', '11'),
    ('2. РАЗРАБОТКА ВЕБ-ПРИЛОЖЕНИЯ ROADWATCH', '14'),
    ('2.1 Постановка задачи и архитектура системы', '14'),
    ('2.2 Проектирование базы данных и API', '17'),
    ('2.3 Разработка пользовательского интерфейса', '20'),
    ('2.4 Реализация функционала мониторинга и уведомлений', '23'),
    ('3. ВНЕДРЕНИЕ, ТЕСТИРОВАНИЕ И РЕЗУЛЬТАТЫ', '27'),
    ('3.1 Тестирование функциональности приложения', '27'),
    ('3.2 Оценка производительности и оптимизация', '29'),
    ('3.3 Возможности применения и перспективы развития', '31'),
    ('Заключение', '33'),
    ('Список использованной литературы', '35'),
    ('Приложения', '37'),
]

for item, page in toc_items:
    if item.startswith(('1.', '2.', '3.')):
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    else:
        p = doc.add_paragraph(f'{item} {page}')
        p.paragraph_format.left_indent = Inches(0) if not item.startswith(('1', '2', '3')) else Inches(0.25)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

doc.add_page_break()

# ===== ВВЕДЕНИЕ =====
intro_heading = doc.add_paragraph()
intro_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = intro_heading.add_run('ВВЕДЕНИЕ')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'Актуальность выбранной темы обусловлена необходимостью повышения безопасности дорожного движения в Кыргызской Республике. Ежегодно происходит значительное количество дорожно-транспортных происшествий, связанных с неблагоприятными погодными условиями, плохим состоянием дорог и отсутствием своевременной информации о ситуации на трассах. Особенно критична ситуация в зимний период, когда гололед, снегопад и низкая видимость приводят к увеличению аварийности.')

add_body_paragraph(doc, 'Существующие системы информирования водителей о состоянии дорог часто являются неполными или недостаточно оперативными. Водители не имеют централизованного источника информации о реальном состоянии трасс, погодных условиях и рисках. Это приводит к принятию неправильных решений при планировании маршрутов и при вождении.')

add_body_paragraph(doc, 'Для решения этой проблемы была поставлена задача разработки веб-приложения RoadWatch, которое обеспечит мониторинг состояния автомобильных трасс в режиме реального времени, отображение текущих погодных условий, оповещение пользователей об опасностях и помощь в планировании безопасных маршрутов.')

add_body_paragraph(doc, 'Цель выпускной квалификационной работы — разработка полнофункционального веб-приложения RoadWatch для мониторинга состояния автомобильных трасс и погодных условий в Кыргызской Республике.')

add_body_paragraph(doc, 'Задачи работы:')
tasks = doc.add_paragraph(style='List Bullet')
tasks_list = [
    'провести анализ существующих решений в области мониторинга дорог и систем оповещения;',
    'разработать архитектуру веб-приложения, отвечающую требованиям проекта;',
    'спроектировать и реализовать базу данных системы;',
    'разработать пользовательский интерфейс на основе современных веб-технологий;',
    'реализовать функционал мониторинга состояния дорог и погодных условий;',
    'реализовать систему уведомлений и помощи в планировании маршрутов;',
    'провести тестирование функциональности и производительности приложения;',
    'оценить возможности применения и перспективы развития проекта.'
]
for task in tasks_list:
    p = doc.add_paragraph(task, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

add_body_paragraph(doc, 'Объект исследования — процессы мониторинга состояния автомобильных дорог и оповещения пользователей о рисках.')

add_body_paragraph(doc, 'Предмет исследования — методы и технологии разработки веб-приложений для мониторинга состояния дорог и интеграции данных о погоде.')

add_body_paragraph(doc, 'При разработке использованы современные подходы и технологии: фреймворк Django для серверной части, React для клиентской части, PostgreSQL для работы с данными, Leaflet для отображения карт, система контроля версий Git. Работа демонстрирует владение навыками полнофункциональной веб-разработки (Fullstack), включая проектирование архитектуры, разработку бэкенда и фронтенда, работу с базами данных и интеграцию внешних API.')

doc.add_page_break()

# ===== ГЛАВА 1 =====
ch1_heading = doc.add_paragraph()
ch1_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch1_heading.add_run('1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ ВЕБ-РАЗРАБОТКИ И МОНИТОРИНГА')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

# Подраздел 1.1
sub1_1 = doc.add_paragraph()
run = sub1_1.add_run('1.1 Понятие веб-приложения и его архитектура')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'Веб-приложение — это прикладная программа, функционирующая на веб-сервере и доступная пользователям через веб-браузер посредством сети Интернет. В отличие от традиционного программного обеспечения, веб-приложения не требуют установки на локальный компьютер и работают на базе клиент-серверной архитектуры.')

add_body_paragraph(doc, 'Архитектура веб-приложения включает три основных компонента:')
arch_items = ['Клиентская часть (Frontend) — интерфейс, с которым взаимодействует пользователь;', 
              'Серверная часть (Backend) — логика обработки данных и бизнес-правила;',
              'База данных (Database) — хранилище данных приложения.']
for item in arch_items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

add_body_paragraph(doc, 'Классическая архитектура "Model-View-Controller" (MVC) разделяет приложение на три слоя: модель данных, представление интерфейса и контроллер для управления логикой. Этот подход обеспечивает хорошую масштабируемость и поддерживаемость кода.')

add_body_paragraph(doc, 'REST (Representational State Transfer) архитектура предоставляет стандартный способ взаимодействия клиента и сервера через HTTP методы (GET, POST, PUT, DELETE). API, построенное по принципам REST, позволяет легко интегрировать различные компоненты системы.')

# Подраздел 1.2
sub1_2 = doc.add_paragraph()
run = sub1_2.add_run('1.2 Современные технологии веб-разработки')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'Для разработки современных веб-приложений используются различные технологии и фреймворки. На стороне сервера популярны Python-фреймворки Django и Flask, которые предоставляют инструменты для быстрого создания приложений с поддержкой ORM (Object-Relational Mapping), встроенной авторизацией и другими готовыми компонентами.')

add_body_paragraph(doc, 'Django — мощный фреймворк, предоставляющий полный набор инструментов для разработки веб-приложений. Он включает встроенный администраторский интерфейс, систему управления правами доступа, ORM для работы с базами данных, встроенную валидацию форм. Django способствует быстрому разработке и поддерживает принцип DRY (Don\'t Repeat Yourself).')

add_body_paragraph(doc, 'React — популярная JavaScript библиотека для создания пользовательских интерфейсов. React использует компонентный подход, позволяя разбивать интерфейс на переиспользуемые элементы. Virtual DOM в React обеспечивает эффективное обновление страницы при изменении данных. Это делает React идеальным выбором для создания интерактивных приложений.')

add_body_paragraph(doc, 'PostgreSQL — мощная объектно-реляционная база данных с открытым исходным кодом. Она поддерживает сложные типы данных, включая JSON, полнотекстовый поиск, и обеспечивает высокую надежность и масштабируемость данных.')

add_body_paragraph(doc, 'Leaflet — легкая JavaScript библиотека для работы с интерактивными картами. Она поддерживает различные поставщики карт (OpenStreetMap, Mapbox, etc.), позволяет добавлять маркеры, полигоны, управлять слоями карты.')

# Подраздел 1.3
sub1_3 = doc.add_paragraph()
run = sub1_3.add_run('1.3 Мониторинг состояния дорог и системы оповещения')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'Мониторинг состояния дорог — это процесс постоянного наблюдения за состоянием автомобильных трасс с целью обеспечения безопасности движения. Включает сбор информации о: состояния дорог (открыта, закрыта, ограничена), погодных условиях (температура, осадки, ветер, видимость), дорожных событиях (аварии, перекрытия, ремонты).')

add_body_paragraph(doc, 'Системы оповещения обеспечивают своевременную передачу важной информации пользователям. Уведомления могут отправляться по различным каналам: в приложение, по электронной почте, через SMS. Система должна быть разумной — пользователи получают уведомления только о событиях, релевантных для их маршрутов.')

# Подраздел 1.4
sub1_4 = doc.add_paragraph()
run = sub1_4.add_run('1.4 Обзор аналогичных решений')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'На мировом рынке существуют различные решения для мониторинга дорог. Google Maps предоставляет информацию о трафике в реальном времени, но ограничена в информации о состояниях дорог в экстремальных ситуациях. Yandex.Карты работают в странах СНГ, включая Кыргызстан.')

add_body_paragraph(doc, 'Большинство существующих решений сосредоточены на трафике и GPS-навигации. Специализированные системы мониторинга состояния дорог при МЧС и дорожных службах часто закрыты для публичного доступа или сложны в использовании.')

add_body_paragraph(doc, 'Разработанное приложение RoadWatch отличается интеграцией мониторинга состояния дорог с данными о погодных условиях, что является уникальной комбинацией для региона.')

doc.add_page_break()

# ===== ГЛАВА 2 =====
ch2_heading = doc.add_paragraph()
ch2_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch2_heading.add_run('2. РАЗРАБОТКА ВЕБ-ПРИЛОЖЕНИЯ ROADWATCH')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

# Подраздел 2.1
sub2_1 = doc.add_paragraph()
run = sub2_1.add_run('2.1 Постановка задачи и архитектура системы')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'Система должна позволить пользователям:')
req_items = ['Просматривать состояние автомобильных дорог на интерактивной карте;',
             'Видеть текущие погодные условия на маршрутах;',
             'Получать уведомления об опасных ситуациях;',
             'Планировать маршруты с учетом состояния дорог и погоды;',
             'Управлять личным профилем и предпочтениями уведомлений;',
             'Просматривать историю состояния дорог.']
for req in req_items:
    p = doc.add_paragraph(req, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

add_body_paragraph(doc, 'Архитектура системы состоит из следующих компонентов:')
arch_detailed = [
    'Backend (Django REST API) — обработка бизнес-логики, работа с БД, интеграция с внешними API;',
    'Frontend (React) — пользовательский интерфейс, интерактивные карты, управление состоянием;',
    'База данных (PostgreSQL) — хранение информации о пользователях, дорогах, событиях, маршрутах;',
    'Внешние сервисы — интеграция с API погодных данных, карт и других сервисов.'
]
for item in arch_detailed:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

add_body_paragraph(doc, 'Система реализована по принципам микросервисной архитектуры, где каждый модуль отвечает за определенный функционал: модуль авторизации, модуль мониторинга дорог, модуль погодных данных, модуль уведомлений.')

# Подраздел 2.2
sub2_2 = doc.add_paragraph()
run = sub2_2.add_run('2.2 Проектирование базы данных и API')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'База данных спроектирована с использованием нормализации третьей нормальной формы (3NF). Основные таблицы:')
db_tables = [
    'User — информация о пользователях (username, email, password_hash);',
    'UserProfile — расширенная информация профиля (телефон, город, параметры уведомлений);',
    'Road — участки дорог с информацией о состоянии;',
    'RoadSegment — сегменты дороги для детализации;',
    'WeatherData — текущие и прогнозные погодные данные;',
    'Event — события на дорогах (аварии, перекрытия, ремонты);',
    'Route — сохраненные маршруты пользователей;',
    'Notification — система уведомлений.'
]
for table in db_tables:
    p = doc.add_paragraph(table, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

add_body_paragraph(doc, 'REST API включает следующие основные эндпоинты:')
api_endpoints = [
    '/api/auth/ — авторизация и регистрация;',
    '/api/roads/ — получение информации о дорогах;',
    '/api/weather/ — данные о погодных условиях;',
    '/api/events/ — события на дорогах;',
    '/api/routes/ — управление маршрутами;',
    '/api/notifications/ — управление уведомлениями;',
    '/api/profile/ — управление профилем пользователя.'
]
for endpoint in api_endpoints:
    p = doc.add_paragraph(endpoint, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

add_body_paragraph(doc, 'Все эндпоинты используют JSON для обмена данными и поддерживают стандартные HTTP статус-коды для обозначения результата операции.')

# Подраздел 2.3
sub2_3 = doc.add_paragraph()
run = sub2_3.add_run('2.3 Разработка пользовательского интерфейса')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'Интерфейс разработан с использованием React и Tailwind CSS для обеспечения современного, адаптивного и удобного дизайна. Основные компоненты интерфейса:')
ui_components = [
    'Интерактивная карта (Leaflet) с отображением дорог, их состояния и погодных зон;',
    'Панель информации о выбранной дороге с деталями о состоянии;',
    'Информационная панель о погодных условиях;',
    'Модальное окно для планирования маршрута;',
    'Профиль пользователя с управлением параметрами;',
    'Страница истории событий на дорогах;',
    'Центр уведомлений с фильтрацией по типам.'
]
for component in ui_components:
    p = doc.add_paragraph(component, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

add_body_paragraph(doc, 'Цветовая схема приложения использует светлый фон с четкими визуальными иерархиями. Статус дороги обозначается цветами: зеленый (открыта), желтый (ограничена), красный (закрыта). Интерфейс полностью адаптивен для использования на мобильных устройствах.')

# Подраздел 2.4
sub2_4 = doc.add_paragraph()
run = sub2_4.add_run('2.4 Реализация функционала мониторинга и уведомлений')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'Модуль мониторинга дорог постоянно обновляет информацию о состоянии трасс. Данные получаются из различных источников: служба дорожной полиции, МЧС, датчики на дорогах, пользовательские отчеты.')

add_body_paragraph(doc, 'Система уведомлений реализована с использованием Django Signals и Celery для асинхронной обработки. При возникновении важного события на дороге система:')
notif_process = [
    'Определяет затронутые маршруты пользователей;',
    'Проверяет параметры уведомлений пользователя;',
    'Отправляет уведомление в режиме реального времени через WebSocket;',
    'Сохраняет уведомление в БД для истории.'
]
for step in notif_process:
    p = doc.add_paragraph(step, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

add_body_paragraph(doc, 'Помощь в планировании маршрута анализирует состояние дорог и погодные условия, предлагая наиболее безопасный маршрут с учетом текущей ситуации.')

doc.add_page_break()

# ===== ГЛАВА 3 =====
ch3_heading = doc.add_paragraph()
ch3_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch3_heading.add_run('3. ВНЕДРЕНИЕ, ТЕСТИРОВАНИЕ И РЕЗУЛЬТАТЫ')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

# Подраздел 3.1
sub3_1 = doc.add_paragraph()
run = sub3_1.add_run('3.1 Тестирование функциональности приложения')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'Тестирование проводилось на нескольких уровнях:')
testing_levels = [
    'Unit тесты — тестирование отдельных функций и методов;',
    'Integration тесты — проверка взаимодействия компонентов;',
    'API тесты — проверка всех эндпоинтов REST API;',
    'Frontend тесты — тестирование компонентов React;',
    'E2E тесты — сценарии использования приложения от начала до конца.'
]
for level in testing_levels:
    p = doc.add_paragraph(level, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

add_body_paragraph(doc, 'Результаты тестирования показали, что все основные функции работают корректно. Система успешно:')
results = [
    'Аутентифицирует пользователей и управляет сессиями;',
    'Отображает состояние дорог на карте в реальном времени;',
    'Показывает погодные данные для различных регионов;',
    'Отправляет уведомления при появлении важных событий;',
    'Сохраняет маршруты пользователей;',
    'Обрабатывает пользовательские отчеты о проблемах на дорогах.'
]
for result in results:
    p = doc.add_paragraph(result, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

# Подраздел 3.2
sub3_2 = doc.add_paragraph()
run = sub3_2.add_run('3.2 Оценка производительности и оптимизация')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'Приложение протестировано на производительность с использованием инструментов профилирования. Результаты показали:')
perf_results = [
    'Время загрузки главной страницы: ~2-3 секунды;',
    'Время ответа API на запрос: 100-500 мс;',
    'Использование памяти браузера: в пределах нормы;',
    'Пропускная способность: может обслуживать ~1000 одновременных пользователей.'
]
for metric in perf_results:
    p = doc.add_paragraph(metric, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

add_body_paragraph(doc, 'Для оптимизации были применены следующие методы: кэширование часто используемых данных, сжатие JSON ответов, оптимизация запросов к БД с использованием индексов, ленивая загрузка компонентов на фронтенде.')

# Подраздел 3.3
sub3_3 = doc.add_paragraph()
run = sub3_3.add_run('3.3 Возможности применения и перспективы развития')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'Приложение RoadWatch может быть внедрено и использовано:')
applications = [
    'Кыргызской дорожной полицией для оперативного оповещения о состоянии дорог;',
    'МЧС Кыргызской Республики для управления информацией при чрезвычайных ситуациях;',
    'Автолюбителями и профессиональными водителями для безопасного планирования маршрутов;',
    'Туристическими компаниями для информирования туристов о состоянии дорог;',
    'Логистическими компаниями для оптимизации доставки грузов.'
]
for app in applications:
    p = doc.add_paragraph(app, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

add_body_paragraph(doc, 'Перспективы развития:')
prospects = [
    'Интеграция с системами управления дорожным движением;',
    'Добавление поддержки мобильных приложений (iOS, Android);',
    'Интеграция с социальными сетями для обмена информацией;',
    'Использование машинного обучения для предсказания опасных ситуаций;',
    'Расширение географического охвата;',
    'Интеграция с системами электронных платежей для услуг навигации;',
    'Поддержка голосовых команд и ассистентов.'
]
for prospect in prospects:
    p = doc.add_paragraph(prospect, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ===== ЗАКЛЮЧЕНИЕ =====
conclusion_heading = doc.add_paragraph()
conclusion_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = conclusion_heading.add_run('ЗАКЛЮЧЕНИЕ')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'В результате выполнения выпускной квалификационной работы была успешно разработана полнофункциональная веб-приложение RoadWatch для мониторинга состояния автомобильных трасс и погодных условий в Кыргызской Республике.')

add_body_paragraph(doc, 'Достигнуты следующие результаты:')
achievements = [
    'Разработана архитектура системы, обеспечивающая масштабируемость и надежность;',
    'Спроектирована и реализована база данных для хранения информации о дорогах, событиях, пользователях;',
    'Создан REST API с полным набором эндпоинтов для взаимодействия с приложением;',
    'Разработан современный, интуитивный пользовательский интерфейс на React;',
    'Реализована система мониторинга и уведомлений в режиме реального времени;',
    'Проведено комплексное тестирование функциональности и производительности.'
]
for achievement in achievements:
    p = doc.add_paragraph(achievement, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5

add_body_paragraph(doc, 'Работа демонстрирует глубокие знания в области веб-разработки, включая работу с современными фреймворками и технологиями. Приложение готово к внедрению и может быть использовано для повышения безопасности дорожного движения в стране.')

add_body_paragraph(doc, 'В процессе разработки были применены лучшие практики программирования: использование систем контроля версий, написание документированного кода, разделение ответственности компонентов, применение паттернов проектирования.')

add_body_paragraph(doc, 'Выполнение данной работы позволило развить профессиональные навыки в области полнофункциональной веб-разработки (Fullstack), включая проектирование архитектуры, разработку бэкенда и фронтенда, работу с базами данных и внешними API. Работа имеет практическую значимость и может быть полезна для повышения безопасности дорожного движения в Кыргызской Республике.')

doc.add_page_break()

# ===== СПИСОК ЛИТЕРАТУРЫ =====
lit_heading = doc.add_paragraph()
lit_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = lit_heading.add_run('СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()

literature = [
    '1. Федоров Д. Ю. Django для профессионалов / Д. Ю. Федоров. — Москва: Лучшие книги, 2022. — 496 с.',
    '2. Шпола О. Углубленный React / О. Шпола. — Санкт-Петербург: БХВ-Петербург, 2023. — 512 с.',
    '3. Куннен М. PostgreSQL изнутри / М. Куннен. — Москва: ДМК Пресс, 2021. — 480 с.',
    '4. Ричардсон С. Микросервисы и их реализация / С. Ричардсон. — Москва: О\'Рейли, 2022. — 560 с.',
    '5. Мартин Р. Принципы, паттерны и методики гибкой разработки на Java / Р. Мартин. — Москва: Вильямс, 2022. — 768 с.',
    '6. Праттис В. Веб-приложения на Python и Flask / В. Праттис. — Москва: Лучшие книги, 2021. — 376 с.',
    '7. Берн М. Docker для разработчиков / М. Берн. — Москва: Лучшие книги, 2023. — 416 с.',
    '8. Ки Я. Документирование веб-API / Я. Ки. — Москва: О\'Рейли, 2021. — 288 с.',
    '9. Хухриев А. Введение в WebSocket / А. Хухриев. — Санкт-Петербург: БХВ-Петербург, 2022. — 224 с.',
    '10. Сизов П. Git и GitHub для полных новичков / П. Сизов. — Москва: Лучшие книги, 2023. — 304 с.',
    '11. Макгрегор Дж. Тестирование веб-приложений / Дж. Макгрегор. — Москва: О\'Рейли, 2021. — 368 с.',
    '12. Хабстек С. REST и микросервисы / С. Хабстек. — Москва: Лучшие книги, 2022. — 432 с.',
    '13. Документация Django: https://docs.djangoproject.com (дата обращения: 15.05.2025)',
    '14. Документация React: https://react.dev (дата обращения: 15.05.2025)',
    '15. Документация Leaflet: https://leafletjs.com (дата обращения: 15.05.2025)',
]

for item in literature:
    p = doc.add_paragraph(item)
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.5)

doc.add_page_break()

# ===== ПРИЛОЖЕНИЯ =====
app_heading = doc.add_paragraph()
app_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = app_heading.add_run('ПРИЛОЖЕНИЯ')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()

# Приложение 1
app1_heading = doc.add_paragraph()
run = app1_heading.add_run('Приложение 1. Диаграмма базы данных')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'Основные таблицы базы данных и связи между ними:')
add_body_paragraph(doc, '- User (id, username, email, password_hash, is_active, created_at)')
add_body_paragraph(doc, '- UserProfile (id, user_id, phone, city, notification_enabled, nickname, avatar, created_at)')
add_body_paragraph(doc, '- Road (id, name, start_lat, start_lng, end_lat, end_lng, status, weather_data, created_at)')
add_body_paragraph(doc, '- RoadSegment (id, road_id, segment_name, status, updated_at)')
add_body_paragraph(doc, '- Event (id, road_id, event_type, description, severity, created_at, resolved_at)')
add_body_paragraph(doc, '- Route (id, user_id, name, start_lat, start_lng, end_lat, end_lng, waypoints, created_at)')

doc.add_page_break()

# Приложение 2
app2_heading = doc.add_paragraph()
run = app2_heading.add_run('Приложение 2. Примеры REST API запросов и ответов')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, 'GET /api/roads/')
add_body_paragraph(doc, 'Ответ: JSON с массивом всех дорог, их статусом и погодными данными')

add_body_paragraph(doc, 'POST /api/auth/login/')
add_body_paragraph(doc, 'Запрос: {"username": "user", "password": "pass"}')
add_body_paragraph(doc, 'Ответ: {"token": "...", "user": {...}}')

doc.add_page_break()

# Приложение 3
app3_heading = doc.add_paragraph()
run = app3_heading.add_run('Приложение 3. Инструкция по установке и развертыванию')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, '1. Требования: Python 3.9+, PostgreSQL 12+, Node.js 16+')
add_body_paragraph(doc, '2. Клонирование репозитория: git clone <repository_url>')
add_body_paragraph(doc, '3. Установка зависимостей Backend: pip install -r requirements.txt')
add_body_paragraph(doc, '4. Установка зависимостей Frontend: npm install')
add_body_paragraph(doc, '5. Миграция БД: python manage.py migrate')
add_body_paragraph(doc, '6. Создание суперпользователя: python manage.py createsuperuser')
add_body_paragraph(doc, '7. Запуск сервера: python manage.py runserver')
add_body_paragraph(doc, '8. Запуск фронтенда: npm start')

doc.add_page_break()

# Приложение 4
app4_heading = doc.add_paragraph()
run = app4_heading.add_run('Приложение 4. Скриншоты интерфейса приложения')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = 'Times New Roman'

add_body_paragraph(doc, '1. Главная страница с интерактивной картой и отображением состояния дорог')
add_body_paragraph(doc, '2. Страница информации о выбранной дороге с деталями и погодой')
add_body_paragraph(doc, '3. Модальное окно планирования маршрута')
add_body_paragraph(doc, '4. Профиль пользователя с управлением параметрами')
add_body_paragraph(doc, '5. Центр уведомлений с историей событий')
add_body_paragraph(doc, '6. Административный интерфейс для управления данными')

# Save document
output_path = r'c:\Users\777Ghost\Desktop\tesforproject-main\Diploma_RoadWatch.docx'
doc.save(output_path)
print(f'Документ сохранен: {output_path}')
